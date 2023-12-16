from django.shortcuts import render, redirect
from .models import Attendance, Attendees, CreateAttendance
from .forms import CreateAttendanceForm
from django.contrib import messages

def createAttendance(request):
    if request.user.is_staff:
        attendance = CreateAttendance.objects.all().order_by('-created_date')
        
        if request.method == "POST":
            form = CreateAttendanceForm(request.POST)
            if form.is_valid():
                form.save()
        else:
            form = CreateAttendanceForm()
        context = {
            'attendance':attendance,
            'form':form,
            }
        return render(request, 'attendance/index.html', context)
    else:
        return redirect('/')
from django.contrib.auth.decorators import login_required
@login_required
def attendanceDetail(request, pk):
    attendance = CreateAttendance.objects.get(pk=pk)
    all = Attendance.objects.filter(attendance=attendance)
  
    context =  {
             'attendance':attendance,
             'attendance_list': {
                 'all': all,
             }
          }
    return render(request, 'attendance/detail.html', context)

from django.http import JsonResponse

def takeAttendance(request, pk, uid):
    attendance = CreateAttendance.objects.get(pk=pk)
    attendee = Attendees.objects.get(uid=uid)
    attend, created = Attendance.objects.get_or_create(attendance=attendance, attendee=attendee)
    data = {
        'attendee':f'{attend.attendee}',
        'time_in':attend.time_in,
    }
    
    return JsonResponse(data, safe=False, status=200)



from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def check_male_gender(lists):
    Male =[]
    for person in lists:
        if person.attendee.Gender == "Male":
            Male.append(person)
    return Male

def check_female_gender(lists):
    Female =[]
    for person in lists:
        if person.attendee.Gender == "Female":
            Female.append(person)
    return Female

def add_title( doc, title_text, size=32, bold=True):
    title = doc.add_paragraph()
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(size)
    if bold:
        title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return title

def generate_word_doc(request, pk):
    attendance = CreateAttendance.objects.get(pk=pk)
    lists = Attendance.objects.filter(attendance=attendance)
    All= []
    for attendee in lists:
        All.append(attendee)
    doc = Document()


    # Add a cover page
    cover_page = doc.sections[0].footer
    cover_page.is_linked_to_previous = False
    cover_page.paragraphs[0].clear()

    title = add_title(doc, f"Again and Afresh Attendance Report \n Gathering of Believers", size=24, bold=True)
    event_name = doc.add_paragraph()
    event_name.add_run(f"Service Title: {attendance.title}").bold = True
    event_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event_date = doc.add_paragraph()
    event_date.add_run(f"Service Date: {attendance.date.date()}").bold = True
    event_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    summaryTableHeaderText = doc.add_paragraph()
    summaryTableHeaderText.add_run(f"Summary Table").bold = True
    summaryTable = doc.add_table(rows=5, cols=5)
    summaryTable.style = 'Table Grid'
    for row in summaryTable.rows:
        for cell in row.cells:
            cell.width = Pt(100)
            cell.paragraphs[0].alignment = 1
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].width = Pt(30)

    summaryTableHeader = summaryTable.rows[0].cells
    summaryTableHeader[0].text = "S/N"
    summaryTableHeader[1].text = "Group"
    summaryTableHeader[2].text = "Male"
    summaryTableHeader[3].text = "Female"
    summaryTableHeader[4].text = "Total"

    summaryTableData = [
        ['1','All', str(len(check_male_gender(All))), str(len(check_female_gender(All))), str(len(All))],
    ]

    for i, row in enumerate(summaryTable.rows[1:], start=1):
        if i - 1 < len(summaryTableData):  # Check if the index is within the range of summaryTableData
            for j, cell in enumerate(row.cells):
                if j < len(summaryTableData[i-1]):  # Check if the index is within the range of summaryTableData[i-1]
                    cell.text = summaryTableData[i-1][j]

    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    ds = len(All) + 2

    add_title(doc, f"Al Attendeesl", size=18, bold=True)
    AllTable = doc.add_table(rows=ds, cols=4)
    AllTable.style = 'Table Grid'

    for row in AllTable.rows:
        for cell in row.cells:
            cell.width = Pt(100)
            cell.paragraphs[0].alignment = 1
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].width = Pt(30)

    AllTableHeader = AllTable.rows[0].cells
    AllTableHeader[0].text = "S/N"
    AllTableHeader[1].text = "Name"
    AllTableHeader[2].text = "Phone Number"
    AllTableHeader[3].text = "Time"

    AllTableData =[]
    for index, member in enumerate(All):
        AllTableData.append(
            [str(index + 1), f'{member.attendee.Full_Name.title()}', str(member.attendee.Phone), str(member.time_in.time())]
        )
    for i, row in enumerate(AllTable.rows[1:], start=1):
        if i - 1 < len(AllTableData):  # Check if the index is within the range of AllTableData
            for j, cell in enumerate(row.cells):
                if j < len(AllTableData[i-1]):  # Check if the index is within the range of AllTableData[i-1]
                    cell.text = AllTableData[i-1][j]
    doc.save(f"Attendance_list/Attendance {attendance.id} for {attendance.title} {attendance.date.date()}.docx")
    print("save")
    messages.success(request, f"You have successfull generated Attendance {attendance.id} for {attendance.title} {attendance.date.date()}.docx")
    return redirect("attendance-list")